from typing import Dict, Any
from .base_generator import BaseFileGenerator

class DocxGenerator(BaseFileGenerator):
    def generate(self, specification: Dict[str, Any], output_path: str, options: Dict[str, Any]) -> None:
        """Generate DOCX file from specification."""
        self.validate_specification(specification)
        
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            raise ImportError("python-docx is required for DOCX generation. Install with: pip install python-docx")
        
        doc = Document()
        
        # Title
        title = doc.add_heading(specification['title'], 0)
        
        # Metadata
        if 'metadata' in specification:
            doc.add_heading('Metadata', level=1)
            for key, value in specification['metadata'].items():
                p = doc.add_paragraph()
                p.add_run(f"{key}: ").bold = True
                p.add_run(str(value))
        
        # Content
        content = specification['content']
        if isinstance(content, str):
            doc.add_paragraph(content)
        elif isinstance(content, list):
            for item in content:
                if isinstance(item, dict):
                    if 'heading' in item:
                        level = item.get('level', 1)
                        doc.add_heading(item['heading'], level=level)
                    if 'text' in item:
                        doc.add_paragraph(item['text'])
                else:
                    doc.add_paragraph(str(item))
        
        # Sections
        if 'sections' in specification:
            for section in specification['sections']:
                doc.add_heading(section.get('title', 'Section'), level=1)
                doc.add_paragraph(section.get('content', ''))
        
        doc.save(output_path)
    
    def get_file_extension(self) -> str:
        return "docx"