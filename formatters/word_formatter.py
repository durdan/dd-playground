from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models.specification import Specification
from formatters.base_formatter import BaseFormatter

class WordFormatter(BaseFormatter):
    def format(self, spec: Specification, output_path: str) -> str:
        """Generate Word document format"""
        doc = Document()
        
        # Set up styles with branding
        self._setup_styles(doc)
        
        # Title page
        title = doc.add_heading(spec.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph()
        metadata_table = doc.add_table(rows=3, cols=2)
        metadata_table.style = 'Table Grid'
        
        cells = metadata_table.rows[0].cells
        cells[0].text = 'Company:'
        cells[1].text = self.branding.company_name
        
        cells = metadata_table.rows[1].cells
        cells[0].text = 'Version:'
        cells[1].text = spec.version
        
        cells = metadata_table.rows[2].cells
        cells[0].text = 'Created:'
        cells[1].text = spec.created_at.strftime('%Y-%m-%d %H:%M:%S')
        
        doc.add_page_break()
        
        # Table of contents
        doc.add_heading('Table of Contents', level=1)
        for i, section in enumerate(spec.sections, 1):
            indent = '\t' * (section.level - 1)
            toc_para = doc.add_paragraph(f"{indent}{i}. {section.title}")
            toc_para.style = 'List Number'
        
        doc.add_page_break()
        
        # Content sections
        for section in spec.sections:
            doc.add_heading(section.title, level=section.level)
            
            # Process content paragraphs
            paragraphs = section.content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph(self.branding.footer_text)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.save(output_path)
        return output_path
    
    def _setup_styles(self, doc):
        """Configure document styles with branding"""
        styles = doc.styles
        
        # Customize heading styles
        heading1 = styles['Heading 1']
        heading1.font.color.rgb = RGBColor.from_string(self.branding.primary_color[1:])
        
        heading2 = styles['Heading 2'] 
        heading2.font.color.rgb = RGBColor.from_string(self.branding.secondary_color[1:])